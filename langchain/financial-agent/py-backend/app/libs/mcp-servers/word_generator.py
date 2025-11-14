from typing import List, Dict, Any, Optional
from mcp.server.fastmcp import FastMCP
import logging
import argparse
import os
import time
import platform
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import json as json_lib

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("word-generator-server")

# Initialize FastMCP server
mcp = FastMCP("word-generator", log_level="INFO")

# Rate limiting implementation
RATE_LIMIT = {
    "per_minute": 20,
    "per_day": 500
}

request_count = {
    "minute": 0,
    "day": 0,
    "last_minute_reset": time.time(),
    "last_day_reset": time.time()
}

def check_rate_limit():
    now = time.time()
    
    if now - request_count["last_minute_reset"] > 60:
        request_count["minute"] = 0
        request_count["last_minute_reset"] = now
    
    if now - request_count["last_day_reset"] > 86400:
        request_count["day"] = 0
        request_count["last_day_reset"] = now
    
    if (request_count["minute"] >= RATE_LIMIT["per_minute"] or 
            request_count["day"] >= RATE_LIMIT["per_day"]):
        raise Exception('Rate limit exceeded. Please try again later.')
    
    request_count["minute"] += 1
    request_count["day"] += 1

class APIError(Exception):
    pass

# Utility functions
def check_file_writeable(filepath: str):
    """
    Check if a file can be written to.
    
    Args:
        filepath: Path to the file
        
    Returns:
        Tuple of (is_writeable, error_message)
    """
    # If file doesn't exist, check if directory is writeable
    if not os.path.exists(filepath):
        directory = os.path.dirname(filepath)
        # If no directory is specified (empty string), use current directory
        if directory == '':
            directory = '.'
        if not os.path.exists(directory):
            return False, f"Directory {directory} does not exist"
        if not os.access(directory, os.W_OK):
            return False, f"Directory {directory} is not writeable"
        return True, ""
    
    # If file exists, check if it's writeable
    if not os.access(filepath, os.W_OK):
        return False, f"File {filepath} is not writeable (permission denied)"
    
    # Try to open the file for writing to see if it's locked
    try:
        with open(filepath, 'a'):
            pass
        return True, ""
    except IOError as e:
        return False, f"File {filepath} is not writeable: {str(e)}"
    except Exception as e:
        return False, f"Unknown error checking file permissions: {str(e)}"

def ensure_docx_extension(filename: str) -> str:
    """
    Ensure filename has .docx extension.
    
    Args:
        filename: The filename to check
        
    Returns:
        Filename with .docx extension
    """
    if not filename.endswith('.docx'):
        return filename + '.docx'
    return filename

def generate_chart_image(chart_data: Dict[str, Any]) -> io.BytesIO:
    """
    Generate a chart image from chart data and return as BytesIO stream.
    
    Args:
        chart_data: Chart configuration and data
        
    Returns:
        BytesIO stream containing the chart image
    """
    try:
        chart_type = chart_data.get('chartType', 'line')
        data = chart_data.get('data', [])
        config = chart_data.get('config', {})
        
        if not data:
            return None
        
        # Create figure and axis
        plt.figure(figsize=(10, 6))
        plt.style.use('default')
        
        if chart_type == 'line':
            # Extract data for line chart
            x_key = config.get('xAxisKey', 'x')
            y_keys = [key for key in chart_data.get('chartConfig', {}).keys()]
            
            for y_key in y_keys:
                x_vals = [item.get(x_key, '') for item in data]
                y_vals = [item.get(y_key, 0) for item in data]
                plt.plot(x_vals, y_vals, marker='o', label=y_key)
            
            plt.legend()
            
        elif chart_type == 'bar':
            # Extract data for bar chart
            x_key = config.get('xAxisKey', 'x')
            y_keys = [key for key in chart_data.get('chartConfig', {}).keys()]
            
            x_vals = [item.get(x_key, '') for item in data]
            
            if len(y_keys) == 1:
                # Single series bar chart
                y_vals = [item.get(y_keys[0], 0) for item in data]
                plt.bar(x_vals, y_vals)
            else:
                # Multiple series bar chart
                import numpy as np
                x_pos = np.arange(len(x_vals))
                width = 0.8 / len(y_keys)
                
                for i, y_key in enumerate(y_keys):
                    y_vals = [item.get(y_key, 0) for item in data]
                    plt.bar(x_pos + i * width, y_vals, width, label=y_key)
                
                plt.xticks(x_pos + width * (len(y_keys) - 1) / 2, x_vals)
                plt.legend()
                
        elif chart_type == 'pie':
            # Extract data for pie chart
            labels = [item.get('segment', item.get('name', '')) for item in data]
            values = [item.get('value', 0) for item in data]
            
            plt.pie(values, labels=labels, autopct='%1.1f%%')
            
        elif chart_type == 'area':
            # Extract data for area chart
            x_key = config.get('xAxisKey', 'x')
            y_keys = [key for key in chart_data.get('chartConfig', {}).keys()]
            
            x_vals = [item.get(x_key, '') for item in data]
            
            for y_key in y_keys:
                y_vals = [item.get(y_key, 0) for item in data]
                plt.fill_between(x_vals, y_vals, alpha=0.7, label=y_key)
            
            plt.legend()
        
        # Set title and labels
        title = config.get('title', 'Chart')
        plt.title(title, fontsize=16, fontweight='bold')
        
        x_label = config.get('xAxisLabel', '')
        y_label = config.get('yAxisLabel', '')
        if x_label:
            plt.xlabel(x_label)
        if y_label:
            plt.ylabel(y_label)
        
        # Improve layout and styling
        plt.tight_layout()
        plt.grid(True, alpha=0.3)
        
        # Save to BytesIO
        img_stream = io.BytesIO()
        plt.savefig(img_stream, format='png', dpi=300, bbox_inches='tight')
        img_stream.seek(0)
        
        # Close the plot to free memory
        plt.close()
        
        return img_stream
        
    except Exception as e:
        logger.error(f"Error generating chart image: {str(e)}")
        plt.close()  # Ensure plot is closed even on error
        return None

async def create_word_document(filename: str, content: str, title: Optional[str] = None, author: Optional[str] = None):
    """
    Create a Word document with formatted content.
    
    Args:
        filename: Path to save the Word document
        content: Text content to include in the document
        title: Optional document title metadata
        author: Optional document author metadata
        
    Returns:
        Status message with result and download info
    """
    check_rate_limit()
    
    try:
        filename = ensure_docx_extension(filename)
        
        # If filename is just a name (no path), save to downloads directory
        if not os.path.dirname(filename):
            # Get downloads directory path - go up 4 levels to reach py-backend, then add downloads
            current_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
            downloads_dir = os.path.join(current_dir, "downloads")
            os.makedirs(downloads_dir, exist_ok=True)
            full_path = os.path.join(downloads_dir, filename)
        else:
            full_path = filename
        
        # Check if file is writeable
        is_writeable, error_message = check_file_writeable(full_path)
        if not is_writeable:
            return f"Cannot create document: {error_message}"
        
        # Create the document
        doc = Document()
        
        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # Process the content and add it to the document
        for element in parse_text_content(content):
            add_element_to_document(doc, element)
        
        # Save the document
        doc.save(full_path)
        
        # Return success message with download info
        just_filename = os.path.basename(full_path)
        return f"Document {just_filename} created successfully. Download available at: /api/files/download/{just_filename}"
    except Exception as e:
        logger.error(f"Error creating document: {str(e)}")
        raise APIError(f"Document creation failed: {str(e)}")

def parse_text_content(content: str) -> List[Dict[str, Any]]:
    """
    Parse plain text into structured elements for a Word document.
    
    Args:
        content: Text content to parse
        
    Returns:
        List of element dictionaries representing document structure
    """
    elements = []
    lines = content.split('\n')
    current_table = None
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Check for headings (lines starting with # characters)
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            elements.append({
                'type': 'heading',
                'text': text,
                'level': level
            })
            i += 1
            continue
        
        # Check for JSON chart data (```json ... ```)
        if line.startswith('```json'):
            # Start collecting JSON content
            json_content = []
            i += 1
            while i < len(lines) and not lines[i].strip() == '```':
                json_content.append(lines[i])
                i += 1
            
            if i < len(lines):  # Found closing ```
                try:
                    json_str = '\n'.join(json_content)
                    chart_data = json_lib.loads(json_str)
                    elements.append({
                        'type': 'chart',
                        'data': chart_data
                    })
                except json_lib.JSONDecodeError:
                    # If JSON is invalid, treat as code block
                    elements.append({
                        'type': 'code',
                        'text': '\n'.join(['```json'] + json_content + ['```'])
                    })
            i += 1
            continue
        
        # Check for table markers (| column | column |)
        if line.startswith('|') and line.endswith('|'):
            # Start of a new table or continuation of an existing one
            if current_table is None:
                current_table = {
                    'type': 'table',
                    'rows': []
                }
                elements.append(current_table)
            
            # Extract cells from the row
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            current_table['rows'].append(cells)
            
            # Check if next line is a separator line (|---|---|)
            if i+1 < len(lines) and re.match(r'^\|[\s\-:]+\|$', lines[i+1].strip()):
                # Skip the separator row
                i += 2
            else:
                i += 1
            continue
        
        # If we reach here and were processing a table, end the table
        if current_table is not None:
            current_table = None
        
        # Check for bullets (lines starting with * or -)
        if line.startswith('* ') or line.startswith('- '):
            bullet_text = line[2:]
            elements.append({
                'type': 'bullet',
                'text': bullet_text
            })
            i += 1
            continue
        
        # Default: treat as regular paragraph
        elements.append({
            'type': 'paragraph',
            'text': line
        })
        i += 1
    
    return elements

def add_element_to_document(doc: Document, element: Dict[str, Any]):
    """
    Add an element to the document based on its type.
    
    Args:
        doc: The document object
        element: Element dictionary from parse_text_content
    """
    element_type = element.get('type', 'paragraph')
    
    if element_type == 'heading':
        doc.add_heading(element['text'], level=element['level'])
    
    elif element_type == 'paragraph':
        doc.add_paragraph(element['text'])
    
    elif element_type == 'bullet':
        paragraph = doc.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(element['text'])
    
    elif element_type == 'chart':
        # Generate chart image and insert into document
        chart_data = element.get('data', {})
        chart_image_stream = generate_chart_image(chart_data)
        if chart_image_stream:
            # Add chart title if available
            chart_title = chart_data.get('config', {}).get('title', 'Chart')
            if chart_title:
                title_para = doc.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_run = title_para.add_run(chart_title)
                title_run.bold = True
                title_run.font.size = Pt(14)
            
            # Insert the chart image
            doc.add_picture(chart_image_stream, width=Inches(6))
            
            # Add some spacing after the chart
            doc.add_paragraph()
    
    elif element_type == 'code':
        # Handle code blocks
        paragraph = doc.add_paragraph()
        paragraph.style = 'Normal'
        run = paragraph.add_run(element['text'])
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    
    elif element_type == 'table':
        rows = element.get('rows', [])
        if rows:
            # Create table with appropriate dimensions
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'
            
            # Fill the table with data
            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    if j < len(table.rows[i].cells):  # Ensure we don't go out of bounds
                        table.rows[i].cells[j].text = cell_text
            
            # Format first row as header if it has content
            if len(rows) > 1:
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

async def parse_table_from_text(text: str) -> List[List[str]]:
    """
    Parse tabular data from text.
    
    Args:
        text: Text containing tabular data
        
    Returns:
        2D list of rows and columns
    """
    check_rate_limit()
    
    try:
        rows = []
        lines = text.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if line:
                # Try to split by common delimiter patterns
                if '|' in line:
                    # Markdown table format: | col1 | col2 |
                    cells = [cell.strip() for cell in line.strip('|').split('|')]
                elif '\t' in line:
                    # Tab-separated values
                    cells = [cell.strip() for cell in line.split('\t')]
                elif ',' in line:
                    # CSV format
                    cells = [cell.strip() for cell in line.split(',')]
                else:
                    # Use multiple spaces as delimiter for fixed-width tables
                    cells = [cell for cell in re.split(r'\s{2,}', line.strip()) if cell]
                
                rows.append(cells)
        
        # Normalize the table to ensure all rows have the same number of columns
        if rows:
            max_cols = max(len(row) for row in rows)
            for row in rows:
                while len(row) < max_cols:
                    row.append('')
        
        return rows
    except Exception as e:
        logger.error(f"Error parsing table: {str(e)}")
        raise APIError(f"Table parsing failed: {str(e)}")

# MCP tool definitions
@mcp.tool()
async def generate_document(filename: str, content: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """
    Generate a formatted Word document from text content.
    
    The content text will be automatically formatted with:
    - Lines starting with # as headings (# to ###### for heading levels 1-6)
    - Lines starting with * or - as bullet points
    - Lines with | column | format as tables
    - Regular text as paragraphs
    
    Args:
        filename: Path to save the Word document (.docx extension optional)
        content: Text content to format into a document
        title: Optional document title for metadata
        author: Optional document author for metadata
    """
    try:
        result = await create_word_document(filename, content, title, author)
        return result
    except Exception as e:
        return f"Error generating document: {str(e)}"

@mcp.tool()
async def create_table_from_text(filename: str, table_text: str, title: Optional[str] = None) -> str:
    """
    Create a Word document with a table parsed from text.
    
    The table text can be in various formats:
    - Markdown tables: | Column1 | Column2 |
    - Tab-separated values
    - Comma-separated values (CSV)
    - Fixed-width columns (separated by 2+ spaces)
    
    Args:
        filename: Path to save the Word document (.docx extension optional)
        table_text: Text containing tabular data
        title: Optional title to add above the table
    """
    try:
        filename = ensure_docx_extension(filename)
        
        # Check if file is writeable
        is_writeable, error_message = check_file_writeable(filename)
        if not is_writeable:
            return f"Cannot create document: {error_message}"
        
        # Parse table structure from text
        table_data = await parse_table_from_text(table_text)
        
        if not table_data:
            return "Could not parse any tabular data from the input text"
        
        # Create document and add title if provided
        doc = Document()
        
        if title:
            doc.add_heading(title, level=1)
        
        # Add table to document
        rows = len(table_data)
        cols = len(table_data[0]) if rows > 0 else 0
        
        if rows > 0 and cols > 0:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'
            
            # Fill table with data
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    if j < cols:  # Ensure within bounds
                        table.cell(i, j).text = cell_text
            
            # Format first row as header if table has multiple rows
            if rows > 1:
                for cell in table.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        
        # Save document
        doc.save(filename)
        return f"Document with table created successfully at {filename}"
    except Exception as e:
        return f"Error creating table document: {str(e)}"

@mcp.tool()
async def generate_document_with_charts(filename: str, content: str, chart_data: Optional[Dict[str, Any]] = None, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """
    Generate a Word document with text content and embedded charts.
    
    Args:
        filename: Path to save the Word document (.docx extension optional)
        content: Text content to format into a document
        chart_data: Optional chart data to embed as images
        title: Optional document title for metadata
        author: Optional document author for metadata
    """
    try:
        result = await create_word_document(filename, content, title, author)
        return result
    except Exception as e:
        return f"Error generating document with charts: {str(e)}"

@mcp.tool()
async def format_financial_report(filename: str, data: Dict[str, Any], title: str = "Financial Report") -> str:
    """
    Create a professionally formatted financial report Word document.
    
    Args:
        filename: Path to save the Word document (.docx extension optional)
        data: Dictionary containing financial data sections
        title: Report title
        
    Data format example:
    {
        "summary": "Financial performance summary text...",
        "highlights": ["Revenue increased by 15%", "Net profit margin improved to 22%"],
        "metrics": {
            "Revenue": "$12.5M", 
            "Expenses": "$8.2M",
            "Net Income": "$4.3M"
        },
        "quarterly_data": [
            ["Q1", "$2.8M", "$1.9M", "$0.9M"],
            ["Q2", "$3.0M", "$2.0M", "$1.0M"],
            ["Q3", "$3.2M", "$2.1M", "$1.1M"],
            ["Q4", "$3.5M", "$2.2M", "$1.3M"]
        ]
    }
    """
    try:
        filename = ensure_docx_extension(filename)
        
        # Check if file is writeable
        is_writeable, error_message = check_file_writeable(filename)
        if not is_writeable:
            return f"Cannot create document: {error_message}"
        
        # Create document
        doc = Document()
        
        # Add title
        doc.add_heading(title, level=0)
        
        # Add summary section
        if "summary" in data:
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(data["summary"])
        
        # Add highlights section
        if "highlights" in data and isinstance(data["highlights"], list):
            doc.add_heading("Key Highlights", level=1)
            for item in data["highlights"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(item)
        
        # Add metrics section
        if "metrics" in data and isinstance(data["metrics"], dict):
            doc.add_heading("Financial Metrics", level=1)
            metrics = data["metrics"]
            table = doc.add_table(rows=len(metrics)+1, cols=2)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Metric"
            header_cells[1].text = "Value"
            
            # Style header row
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add metric rows
            i = 1
            for metric_name, metric_value in metrics.items():
                if i < len(table.rows):
                    row = table.rows[i]
                    row.cells[0].text = str(metric_name)
                    row.cells[1].text = str(metric_value)
                    i += 1
        
        # Add quarterly data section
        if "quarterly_data" in data and isinstance(data["quarterly_data"], list):
            quarterly_data = data["quarterly_data"]
            if quarterly_data:
                doc.add_heading("Quarterly Performance", level=1)
                
                # Add header row if not already included
                header_row = ["Quarter", "Revenue", "Expenses", "Net Income"]
                if len(quarterly_data) > 0 and "Q" not in str(quarterly_data[0][0]):
                    quarterly_data.insert(0, header_row)
                
                # Create table
                rows = len(quarterly_data)
                cols = len(quarterly_data[0]) if rows > 0 else 0
                
                if rows > 0 and cols > 0:
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    
                    # Fill table with data
                    for i, row_data in enumerate(quarterly_data):
                        for j, cell_text in enumerate(row_data):
                            if j < cols:  # Ensure within bounds
                                table.cell(i, j).text = str(cell_text)
                    
                    # Format header row
                    for cell in table.rows[0].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        
        # Save document
        doc.save(filename)
        return f"Financial report created successfully at {filename}"
    except Exception as e:
        return f"Error creating financial report: {str(e)}"

def main():
    parser = argparse.ArgumentParser(description='Run Word Generator MCP server')
    parser.add_argument('--port', type=int, default=8089, help='Port to run the server on')
    
    args = parser.parse_args()
    
    mcp.settings.port = args.port
    logger.info(f"Starting Word Generator server with Streamable HTTP transport on port {args.port}")
    mcp.run(transport='streamable-http')

if __name__ == "__main__":
    main()